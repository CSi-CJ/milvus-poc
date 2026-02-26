"""
文本文件解析器
"""

import os
from typing import Dict, Any, List
import logging
import io

try:
    import chardet
except ImportError:
    chardet = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

from ..models import ParsedContent
from .base import BaseFileParser


class TextParser(BaseFileParser):
    """文本文件解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.txt', '.md', '.doc', '.docx'}
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否能解析文本文件"""
        ext = self._get_file_extension(file_path)
        return ext in self.supported_extensions
    
    def parse(self, file_path: str) -> ParsedContent:
        """解析文本文件"""
        self._validate_file(file_path)
        
        ext = self._get_file_extension(file_path)
        
        try:
            if ext in {'.txt', '.md'}:
                return self._parse_plain_text(file_path)
            elif ext == '.docx':
                return self._parse_docx(file_path)
            elif ext == '.doc':
                return self._parse_doc(file_path)
            else:
                return self._create_error_content(file_path, f"Unsupported text file type: {ext}")
                
        except Exception as e:
            self.logger.error(f"Error parsing text file {file_path}: {e}")
            return self._create_error_content(file_path, str(e))
    
    def _parse_plain_text(self, file_path: str) -> ParsedContent:
        """解析纯文本文件"""
        # 检测编码
        encoding = self._detect_encoding(file_path)
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
        except UnicodeDecodeError:
            # 如果检测的编码失败，尝试常见编码
            for fallback_encoding in ['utf-8', 'gbk', 'gb2312', 'latin1']:
                try:
                    with open(file_path, 'r', encoding=fallback_encoding) as f:
                        content = f.read()
                    encoding = fallback_encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise UnicodeDecodeError("Unable to decode file with any common encoding")
        
        metadata = {
            'encoding': encoding,
            'line_count': content.count('\n') + 1,
            'char_count': len(content),
            'word_count': len(content.split()) if content.strip() else 0,
        }
        
        # 生成文本可视化截图
        images = []
        if content.strip():
            try:
                text_image = self._generate_text_screenshot(content, file_path)
                if text_image:
                    images.append(text_image)
            except Exception as e:
                self.logger.warning(f"Failed to generate text screenshot for {file_path}: {e}")
        
        return ParsedContent(
            text_content=content if content.strip() else None,
            image_content=images if images else None,
            audio_content=None,
            metadata=metadata,
            file_type=self._get_file_extension(file_path)
        )
    
    def _parse_docx(self, file_path: str) -> ParsedContent:
        """解析 DOCX 文件"""
        if Document is None:
            return self._create_error_content(
                file_path,
                "python-docx not installed. Please install with: pip install python-docx"
            )
        
        try:
            doc = Document(file_path)
            
            # 提取文本内容
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            content = '\n'.join(paragraphs)
            
            # 提取元数据
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(content.split()) if content.strip() else 0,
            }
            
            # 生成Word文档可视化截图
            images = []
            if content.strip():
                try:
                    docx_image = self._generate_docx_screenshot(content, file_path)
                    if docx_image:
                        images.append(docx_image)
                except Exception as e:
                    self.logger.warning(f"Failed to generate DOCX screenshot for {file_path}: {e}")
            
            return ParsedContent(
                text_content=content if content.strip() else None,
                image_content=images if images else None,
                audio_content=None,
                metadata=metadata,
                file_type='.docx'
            )
            
        except Exception as e:
            self.logger.error(f"Error parsing DOCX file {file_path}: {e}")
            return self._create_error_content(file_path, str(e))
    
    def _parse_doc(self, file_path: str) -> ParsedContent:
        """解析 DOC 文件（旧格式）"""
        # DOC 格式比较复杂，这里提供一个基本实现
        # 在生产环境中，建议使用 python-docx2txt 或转换为 DOCX
        return self._create_error_content(
            file_path,
            "DOC format not fully supported. Please convert to DOCX format."
        )
    
    def _detect_encoding(self, file_path: str) -> str:
        """检测文件编码"""
        if chardet is None:
            return 'utf-8'  # 默认编码
        
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # 读取前10KB用于检测
            
            result = chardet.detect(raw_data)
            encoding = result.get('encoding', 'utf-8')
            confidence = result.get('confidence', 0)
            
            # 如果置信度太低，使用默认编码
            if confidence < 0.7:
                encoding = 'utf-8'
            
            self.logger.debug(f"Detected encoding for {file_path}: {encoding} (confidence: {confidence})")
            return encoding
            
        except Exception as e:
            self.logger.warning(f"Failed to detect encoding for {file_path}: {e}")
            return 'utf-8'
    
    def _generate_text_screenshot(self, content: str, file_path: str) -> bytes:
        """为纯文本生成可视化截图"""
        if not PIL_AVAILABLE:
            self.logger.warning("PIL not available, cannot generate text screenshot")
            return None
        
        try:
            # 限制内容长度以避免图像过大
            max_chars = 2000
            if len(content) > max_chars:
                content = content[:max_chars] + "\n... (内容已截断)"
            
            # 设置图像参数
            width = 800
            font_size = 14
            line_height = font_size + 4
            padding = 20
            
            # 尝试加载字体
            try:
                # Windows系统字体
                font = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", font_size)
            except:
                try:
                    # 备用字体
                    font = ImageFont.truetype("arial.ttf", font_size)
                except:
                    # 默认字体
                    font = ImageFont.load_default()
            
            # 分割文本为行
            lines = content.split('\n')
            
            # 计算图像高度
            height = len(lines) * line_height + padding * 2
            height = min(height, 2000)  # 限制最大高度
            
            # 创建图像
            image = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(image)
            
            # 绘制文本
            y = padding
            for line in lines:
                if y + line_height > height - padding:
                    break
                
                # 处理长行
                if len(line) > 80:
                    line = line[:80] + "..."
                
                draw.text((padding, y), line, fill='black', font=font)
                y += line_height
            
            # 添加文件名标题
            title = f"文件: {os.path.basename(file_path)}"
            draw.text((padding, 5), title, fill='blue', font=font)
            
            # 转换为字节
            buffer = io.BytesIO()
            image.save(buffer, format='PNG')
            return buffer.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error generating text screenshot: {e}")
            return None
    
    def _generate_docx_screenshot(self, content: str, file_path: str) -> bytes:
        """为Word文档生成可视化截图"""
        if not PIL_AVAILABLE:
            self.logger.warning("PIL not available, cannot generate DOCX screenshot")
            return None
        
        try:
            # 限制内容长度
            max_chars = 2000
            if len(content) > max_chars:
                content = content[:max_chars] + "\n... (内容已截断)"
            
            # 设置图像参数（模拟Word文档样式）
            width = 850
            font_size = 12
            line_height = font_size + 6
            padding = 40
            
            # 尝试加载字体
            try:
                font = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", font_size)
                title_font = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", font_size + 2)
            except:
                try:
                    font = ImageFont.truetype("arial.ttf", font_size)
                    title_font = ImageFont.truetype("arial.ttf", font_size + 2)
                except:
                    font = ImageFont.load_default()
                    title_font = font
            
            # 分割文本为段落
            paragraphs = content.split('\n\n')
            
            # 计算图像高度
            total_lines = sum(len(p.split('\n')) for p in paragraphs) + len(paragraphs)
            height = total_lines * line_height + padding * 2 + 50
            height = min(height, 2000)
            
            # 创建图像（模拟Word页面）
            image = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(image)
            
            # 绘制页面边框
            draw.rectangle([10, 10, width-10, height-10], outline='lightgray', width=1)
            
            # 添加文档标题
            title = f"Word文档: {os.path.basename(file_path)}"
            draw.text((padding, 20), title, fill='darkblue', font=title_font)
            
            # 绘制内容
            y = 60
            for paragraph in paragraphs:
                if y + line_height > height - padding:
                    break
                
                lines = paragraph.split('\n')
                for line in lines:
                    if y + line_height > height - padding:
                        break
                    
                    # 处理长行
                    if len(line) > 70:
                        line = line[:70] + "..."
                    
                    draw.text((padding, y), line, fill='black', font=font)
                    y += line_height
                
                y += line_height // 2  # 段落间距
            
            # 转换为字节
            buffer = io.BytesIO()
            image.save(buffer, format='PNG')
            return buffer.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error generating DOCX screenshot: {e}")
            return None